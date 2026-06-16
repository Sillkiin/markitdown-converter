#!/usr/bin/env python3
"""
convert.py - Convert documents/URLs to clean Markdown using Microsoft MarkItDown.

Exposes MarkItDown's full feature set and degrades gracefully: if an optional
backend dependency is missing (or produces empty output because a backend such
as exiftool / a transcription model isn't installed), it falls back to a
best-effort extractor so a conversion still yields useful content instead of an
error or an empty file. The script never emits a raw Python traceback and exits
non-zero only if an input genuinely could not be read.

Usage:
    python convert.py INPUT [INPUT ...] [-o OUTPUT] [options]

Inputs:
  - A file path, a directory (its top-level files are converted; sub-directories
    are skipped with a note), or a URL (e.g. a YouTube link).

Output:
  - One file, no -o:        writes <basename>.md in the current dir AND prints to stdout.
  - One file, -o file.md:   writes to that file.
  - One file, -o dir/:      writes <basename>.md into that directory.
  - Many inputs / a dir:    -o is treated as a directory; one .md per input is
                            written there. Original extension is kept in the name
                            and a counter is appended if two inputs would collide.

Full-functionality options (all optional):
  --use-plugins             Enable installed 3rd-party MarkItDown plugins.
  --llm-model MODEL         Use an LLM for image descriptions (e.g. gpt-4o). Needs `openai`.
  --llm-api-key KEY         API key for the LLM client (else uses env, e.g. OPENAI_API_KEY).
  --llm-base-url URL        Custom base URL for an OpenAI-compatible endpoint.
  --docintel-endpoint URL   Use Azure Document Intelligence for higher-quality extraction.
  --recursive               Recurse into sub-directories instead of skipping them.

Install MarkItDown with every backend:
    pip install 'markitdown[all]'    # add --break-system-packages on a managed Python
"""
import argparse
import codecs
import os
import sys
import subprocess


SUPPORTED = ("PDF, Word (.docx), PowerPoint (.pptx), Excel (.xlsx/.xls), images "
             "(metadata/EXIF + OCR), audio (metadata + transcription), HTML, CSV, TSV, "
             "JSON, XML, RSS, EPub, .ipynb, ZIP, YouTube URLs, and more")

DETECT_BYTES = 262144  # bytes read for charset detection (avoid loading huge files)


def ensure_markitdown():
    """Import MarkItDown, installing 'markitdown[all]' on demand if needed.
    pip output is sent to stderr so it never contaminates the Markdown on stdout."""
    try:
        from markitdown import MarkItDown  # noqa: F401
        return
    except ImportError:
        pass

    sys.stderr.write("markitdown not found - attempting to install 'markitdown[all]'...\n")
    for cmd in (
        [sys.executable, "-m", "pip", "install", "markitdown[all]"],
        [sys.executable, "-m", "pip", "install", "--break-system-packages", "markitdown[all]"],
    ):
        try:
            subprocess.run(cmd, check=True, stdout=sys.stderr, stderr=sys.stderr)
            from markitdown import MarkItDown  # noqa: F401
            return
        except Exception:
            continue

    sys.stderr.write(
        "\nERROR: could not import or install markitdown.\n"
        "Install it manually and re-run:\n"
        "    pip install 'markitdown[all]'   "
        "(add --break-system-packages on a managed Python)\n"
    )
    sys.exit(1)


MIN_MARKITDOWN = (0, 1, 6)  # version this wrapper was validated against


def warn_if_old_markitdown():
    """Soft, non-fatal heads-up if an older MarkItDown is installed (field/method
    names have shifted across versions; the wrapper already handles known cases)."""
    try:
        import markitdown
        v = getattr(markitdown, "__version__", "") or ""
        parts = tuple(int(x) for x in v.split(".")[:3] if x.isdigit())
        if parts and parts < MIN_MARKITDOWN:
            sys.stderr.write("NOTE: markitdown %s is older than the validated %s; "
                             "consider: pip install -U 'markitdown[all]'\n"
                             % (v, ".".join(map(str, MIN_MARKITDOWN))))
    except Exception:
        pass


def build_converter(args):
    """Construct a MarkItDown instance honoring all full-functionality options."""
    from markitdown import MarkItDown
    kwargs = {"enable_plugins": bool(args.use_plugins)}
    if args.docintel_endpoint:
        kwargs["docintel_endpoint"] = args.docintel_endpoint
    if args.llm_model:
        try:
            from openai import OpenAI
        except ImportError:
            sys.stderr.write("ERROR: --llm-model needs the openai package: pip install openai\n")
            sys.exit(1)
        client_kwargs = {}
        if args.llm_api_key:
            client_kwargs["api_key"] = args.llm_api_key
        if args.llm_base_url:
            client_kwargs["base_url"] = args.llm_base_url
        kwargs["llm_client"] = OpenAI(**client_kwargs)
        kwargs["llm_model"] = args.llm_model
    return MarkItDown(**kwargs)


def is_url(s):
    low = s.lower()
    return low.startswith("http://") or low.startswith("https://")


def expand_inputs(inputs, recursive=False):
    """Expand directories into contained files; keep files and URLs as-is.
    By default only top-level files of a directory are taken; with recursive=True
    the whole tree is walked. Skipped sub-directories are reported to stderr."""
    expanded = []
    skipped_dirs = 0
    for item in inputs:
        if is_url(item):
            expanded.append(item)
        elif os.path.isdir(item):
            if recursive:
                for root, _dirs, files in os.walk(item):
                    for name in sorted(files):
                        expanded.append(os.path.join(root, name))
            else:
                for name in sorted(os.listdir(item)):
                    full = os.path.join(item, name)
                    if os.path.isfile(full):
                        expanded.append(full)
                    elif os.path.isdir(full):
                        skipped_dirs += 1
        else:
            expanded.append(item)
    if skipped_dirs:
        sys.stderr.write("NOTE: skipped %d sub-directory(ies); pass --recursive to include them.\n"
                         % skipped_dirs)
    return expanded


_ILLEGAL_FN_CHARS = '<>:"/\\|?*'


def _sanitize_filename(name):
    """Strip characters illegal in a filename and never return ''. Critically, a
    bare ':' on NTFS makes open('name:stream', 'w') succeed by writing to a hidden
    Alternate Data Stream instead of failing - so an unsanitized URL-derived name
    like 'report:final.pdf' silently loses the output to an ADS while reporting
    success. Replace the Windows-illegal set and control chars, drop trailing dots/
    spaces (also illegal on Windows), and reject pure-dot/empty names."""
    cleaned = "".join("_" if (c in _ILLEGAL_FN_CHARS or ord(c) < 32) else c
                      for c in name).strip(" .")
    return cleaned if cleaned not in ("", ".", "..") else "output"


def out_name_for(inp, keep_ext=False):
    """Derive a safe base .md filename for a given input file or URL."""
    if is_url(inp):
        from urllib.parse import urlsplit
        import posixpath
        # urlsplit drops the query/fragment; posixpath.basename takes the last path
        # segment without mis-parsing a query value that contains '/'.
        base = posixpath.basename(urlsplit(inp).path.rstrip("/")) or "output"
        return _sanitize_filename(base) + ".md"
    fname = os.path.basename(inp.rstrip("/\\")) or "output"
    if keep_ext:
        return _sanitize_filename(fname) + ".md"
    return _sanitize_filename(os.path.splitext(fname)[0] or "output") + ".md"


def unique_target(directory, name, used):
    """Return a collision-free path inside *directory* for *name*, remembering
    chosen paths in *used* so two different inputs never silently overwrite."""
    stem, ext = os.path.splitext(name)
    candidate = os.path.join(directory, name)
    i = 1
    while candidate in used or os.path.exists(candidate):
        candidate = os.path.join(directory, "%s-%d%s" % (stem, i, ext))
        i += 1
    used.add(candidate)
    return candidate


# ----------------------------- encoding handling -----------------------------

TEXT_EXTS = {".csv", ".tsv", ".txt", ".text", ".md", ".markdown",
             ".json", ".jsonl", ".xml", ".html", ".htm", ".log", ".rss", ".atom"}


def _looks_utf8(prefix):
    """True if *prefix* is valid UTF-8, tolerating a multibyte char cut off at the
    very end (which happens because we only read a prefix for detection)."""
    try:
        prefix.decode("utf-8")
        return True
    except UnicodeDecodeError as e:
        # Tolerate ONLY a genuine multibyte char truncated by the prefix boundary,
        # not an invalid byte that merely happens to sit near the end.
        if e.reason == "unexpected end of data" and e.start >= len(prefix) - 3:
            try:
                prefix[:e.start].decode("utf-8")
                return True
            except Exception:
                return False
        return False


def _whole_file_decodes(path, enc):
    """True if the ENTIRE file decodes under *enc*, streamed in 1 MB chunks through
    an incremental decoder (O(1) memory, correct across multibyte chunk boundaries).
    detect_charset only reads a bounded prefix, so a file that is pure ASCII for the
    first 256 KB and then has single-byte (cp1251/latin-1) content would otherwise be
    wrongly committed to utf-8 and fail the full-file decode downstream. Used to
    confirm a charset over the whole file before trusting it."""
    try:
        dec = codecs.getincrementaldecoder(enc)()
    except LookupError:
        return False
    try:
        with open(path, "rb") as fh:
            while True:
                chunk = fh.read(1 << 20)
                if not chunk:
                    dec.decode(b"", final=True)
                    return True
                dec.decode(chunk)
    except (UnicodeDecodeError, OSError):
        return False


def _detect_utf16_no_bom(raw):
    """Return 'utf-16-le'/'utf-16-be'/'utf-16' if *raw* looks like BOM-less UTF-16,
    else None. ASCII text encoded as UTF-16 is ~50% NUL bytes (every other byte),
    which is otherwise valid UTF-8 and would be silently mojibake'd. NUL position
    (odd vs even) reveals the byte order."""
    if not raw:
        return None
    if raw.count(0) / len(raw) <= 0.25:
        return None
    even_nul = raw[0::2].count(0)
    odd_nul = raw[1::2].count(0)
    if odd_nul > even_nul * 3:
        return "utf-16-le"   # NUL in the high byte -> little-endian
    if even_nul > odd_nul * 3:
        return "utf-16-be"
    return "utf-16"


# charset_normalizer is unreliable on short, mostly-ASCII byte strings: for a few
# high bytes it frequently returns an arbitrary CJK multibyte codec. These are the
# usual culprits; when one is guessed for predominantly-ASCII data we distrust it.
_CJK_MULTIBYTE = {
    "big5", "big5hkscs", "gb2312", "gbk", "gb18030", "hz",
    "euckr", "eucjp", "eucjis2004", "eucjisx0213", "shiftjis", "shiftjis2004",
    "shiftjisx0213", "cp932", "cp949", "cp950", "johab",
    "iso2022jp", "iso2022jp1", "iso2022jp2", "iso2022jp2004", "iso2022jp3",
    "iso2022jpext", "iso2022kr",
}
_MOSTLY_ASCII_RATIO = 0.30  # below this non-ASCII share, a CJK guess is suspect


def _norm_enc(name):
    return "".join(ch for ch in (name or "").lower() if ch.isalnum())


def _decodes_clean(raw, enc):
    try:
        raw.decode(enc)
        return True
    except (UnicodeDecodeError, LookupError):
        return False


def _locale_single_byte_candidates():
    """Single-byte encodings to try when a CJK guess is distrusted, locale first."""
    cands = []
    try:
        import locale
        enc = locale.getpreferredencoding(False)
        if enc and _norm_enc(enc) not in _CJK_MULTIBYTE and _norm_enc(enc) not in ("utf8", "utf8sig"):
            cands.append(enc)
    except Exception:
        pass
    for enc in ("cp1251", "cp1252"):
        if enc not in cands:
            cands.append(enc)
    return cands


def detect_charset(path):
    """Return a reliable charset, preferring UTF-8 when the bytes are valid UTF-8.
    Only a bounded prefix is read so a multi-GB text file can't exhaust memory."""
    with open(path, "rb") as fh:
        raw = fh.read(DETECT_BYTES)
    if raw.startswith(b"\xef\xbb\xbf"):
        return "utf-8-sig"
    if raw.startswith((b"\xff\xfe", b"\xfe\xff")):
        return "utf-16"  # BOM-tagged UTF-16 (charset_normalizer also handles it)
    # BOM-less UTF-16 is ~50% NUL bytes, which are valid UTF-8 and would otherwise
    # be accepted below and silently produce NUL-interleaved mojibake.
    utf16 = _detect_utf16_no_bom(raw)
    if utf16:
        return utf16
    # Trust utf-8 only if the WHOLE file is valid utf-8, not merely the prefix:
    # a large file whose non-ASCII bytes start past the prefix must not be
    # mis-committed to utf-8 (its full-file decode would then fail).
    if _looks_utf8(raw):
        if len(raw) < DETECT_BYTES or _whole_file_decodes(path, "utf-8"):
            return "utf-8"
        # Prefix is ASCII/utf-8 but the full file is not: the distinguishing
        # single-byte content is past the detection window, so charset_normalizer
        # on the prefix can't see it. Pick a single-byte encoding that decodes the
        # whole file (locale first, then latin-1 which never fails) to avoid both
        # a spurious decode failure and silent mojibake.
        for enc in _locale_single_byte_candidates():
            if _whole_file_decodes(path, enc):
                return enc
        return "latin-1"

    guess = None
    try:
        from charset_normalizer import from_bytes
        best = from_bytes(raw).best()
        if best is not None and best.encoding:
            guess = best.encoding
    except Exception:
        pass

    # Distrust a CJK multibyte guess on predominantly-ASCII data (e.g. a JSON/CSV
    # file with a few cp1251 Cyrillic bytes mis-read as Big5). Prefer the machine's
    # locale single-byte encoding when it decodes the bytes cleanly. Dense CJK text
    # (high non-ASCII ratio) keeps the guess, so real Big5/GBK files still work.
    if guess and _norm_enc(guess) in _CJK_MULTIBYTE:
        nonascii = sum(1 for b in raw if b >= 0x80)
        if raw and nonascii / len(raw) < _MOSTLY_ASCII_RATIO:
            for enc in _locale_single_byte_candidates():
                if _decodes_clean(raw, enc):
                    return enc

    if guess:
        return guess
    for enc in ("cp1251", "latin-1"):
        if _decodes_clean(raw, enc):
            return enc
    return "utf-8"


def _text_stream_info(inp, ext):
    try:
        from markitdown import StreamInfo
    except Exception:
        try:
            from markitdown._stream_info import StreamInfo
        except Exception:
            return None
    try:
        return StreamInfo(extension=ext, charset=detect_charset(inp))
    except Exception:
        return None


# ----------------------------- fallback extractors ---------------------------

def _md_cell(value):
    """Render a value as Markdown table-cell text that cannot break the table:
    escape backslashes and pipes, and turn newlines into <br>. Without this a
    cell containing '|' (URLs, free text) silently adds a column and a cell with
    an embedded newline (Excel Alt+Enter) splits the row across lines."""
    s = str(value)
    s = s.replace("\\", "\\\\").replace("|", "\\|")
    return s.replace("\r\n", "<br>").replace("\n", "<br>").replace("\r", "<br>")


def docx_fallback(path):
    """.docx -> Markdown via python-docx when MarkItDown's mammoth backend is absent."""
    from docx import Document
    doc = Document(path)
    lines = []
    for p in doc.paragraphs:
        t = p.text.strip()
        style = (getattr(p.style, "name", "") or "").lower()
        if not t:
            lines.append("")
            continue
        if style.startswith("heading") or style == "title":
            digits = "".join(ch for ch in style if ch.isdigit())
            lvl = int(digits) if digits else 1
            lines.append("#" * max(1, min(lvl, 6)) + " " + t)
        elif "list" in style:
            lines.append("- " + t)
        else:
            lines.append(t)
    for tbl in doc.tables:
        rows = tbl.rows
        if not rows:
            continue
        header = [_md_cell(c.text.strip()) for c in rows[0].cells]
        lines.append("")
        lines.append("| " + " | ".join(header) + " |")
        lines.append("| " + " | ".join("---" for _ in header) + " |")
        for r in rows[1:]:
            lines.append("| " + " | ".join(_md_cell(c.text.strip()) for c in r.cells) + " |")
    return "\n".join(lines).strip() + "\n"


def xlsx_fallback(path):
    """.xlsx -> Markdown tables via openpyxl, one table per sheet. Cells are
    escaped so pipes/newlines can't break the table. A formula cell with no
    cached value (e.g. a workbook last saved by a non-Excel tool, which never
    computes formulas) falls back to its formula text instead of a silent blank."""
    import openpyxl
    wb = openpyxl.load_workbook(path, data_only=True)
    try:
        formula_sheets = openpyxl.load_workbook(path, data_only=False).worksheets
    except Exception:
        formula_sheets = [None] * len(wb.worksheets)
    out = []
    for ws, ws_f in zip(wb.worksheets, formula_sheets):
        vrows = list(ws.iter_rows(values_only=True))
        frows = list(ws_f.iter_rows(values_only=True)) if ws_f is not None else []
        rows = []
        for ri, vrow in enumerate(vrows):
            frow = frows[ri] if ri < len(frows) else ()
            cells = []
            for ci, v in enumerate(vrow):
                if v is None and ci < len(frow):
                    f = frow[ci]
                    if isinstance(f, str) and f.startswith("="):
                        v = f  # uncached formula: surface the formula text
                cells.append("" if v is None else _md_cell(v))
            rows.append(cells)
        rows = [r for r in rows if any(cell != "" for cell in r)]
        if not rows:
            continue
        out.append("## " + ws.title)
        width = max(len(r) for r in rows)
        rows = [r + [""] * (width - len(r)) for r in rows]
        out.append("| " + " | ".join(rows[0]) + " |")
        out.append("| " + " | ".join("---" for _ in range(width)) + " |")
        for r in rows[1:]:
            out.append("| " + " | ".join(r) + " |")
        out.append("")
    return "\n".join(out).strip() + "\n"


def _sniff_delimiter(header_line):
    """Pick a delimiter from the header line, preferring comma on a tie. MarkItDown's
    native CSV reader does not sniff, so a European 'name;city;score' .csv otherwise
    collapses into a single column; this lets it tabulate without regressing comma
    CSVs. Quoting-aware: it parses the line with each candidate delimiter via the csv
    module and picks the one yielding the most fields, so delimiters that merely sit
    INSIDE a quoted cell (e.g. a comma CSV with a "see a; b; c" notes column) don't
    win - a naive character count would mis-pick ';' and collapse the table."""
    import csv as _csv
    import io
    best, best_n = ",", 0
    for d in (",", ";", "\t"):
        try:
            fields = next(_csv.reader(io.StringIO(header_line), delimiter=d), [])
        except Exception:
            fields = []
        if len(fields) > best_n:
            best, best_n = d, len(fields)
    return best


def csv_fallback(path, delimiter=None):
    """.csv/.tsv -> escaped Markdown table. MarkItDown's native CSV converter does
    NOT escape '|' or embedded newlines, so a cell containing a pipe silently adds a
    phantom column and an Excel-style multiline cell (a quoted newline) splits the
    row across lines - the exact breakage xlsx_fallback already guards against. Parse
    with the csv module (correct RFC-4180 quoting) using the detected charset and
    escape every cell via _md_cell so the table cannot break. Returns '' on empty
    input so convert_one falls through to the native path."""
    import csv as _csv
    import io
    enc = detect_charset(path)
    with open(path, "rb") as fh:
        data = fh.read()
    try:
        text = data.decode(enc)
    except (UnicodeDecodeError, LookupError):
        text = data.decode("latin-1")  # never fails; better than a crash/exit2
    text = text.replace("\x00", "")  # a stray NUL would make csv.reader raise
    if delimiter is None:
        first = next((ln for ln in text.splitlines() if ln.strip()), "")
        delimiter = _sniff_delimiter(first)
    rows = list(_csv.reader(io.StringIO(text), delimiter=delimiter))
    rows = [r for r in rows if any((c or "").strip() != "" for c in r)]
    if not rows:
        return ""
    width = max(len(r) for r in rows)
    rows = [list(r) + [""] * (width - len(r)) for r in rows]
    lines = ["| " + " | ".join(_md_cell(c) for c in rows[0]) + " |",
             "| " + " | ".join("---" for _ in range(width)) + " |"]
    for r in rows[1:]:
        lines.append("| " + " | ".join(_md_cell(c) for c in r) + " |")
    return "\n".join(lines).strip() + "\n"


def image_fallback(path):
    """Image -> Markdown metadata via Pillow: format, dimensions, embedded text, EXIF."""
    from PIL import Image, ExifTags
    name = os.path.basename(path)
    lines = ["# Image: %s" % name, ""]
    with Image.open(path) as im:
        lines.append("- **Format:** %s" % im.format)
        lines.append("- **Dimensions:** %d x %d" % (im.width, im.height))
        lines.append("- **Mode:** %s" % im.mode)
        for k, v in (im.info or {}).items():
            if isinstance(v, str) and k.lower() not in ("exif", "icc_profile"):
                lines.append("- **%s:** %s" % (k, v))
        try:
            exif = im.getexif()
            for tag_id, val in (exif.items() if exif else []):
                tag = ExifTags.TAGS.get(tag_id, str(tag_id))
                if isinstance(val, (str, int, float)):
                    lines.append("- **%s:** %s" % (tag, val))
        except Exception:
            pass
    return "\n".join(lines).strip() + "\n"


def wav_fallback(path):
    """.wav -> Markdown metadata via the stdlib wave module."""
    import wave
    name = os.path.basename(path)
    with wave.open(path, "rb") as w:
        ch, sw, fr, n = w.getnchannels(), w.getsampwidth(), w.getframerate(), w.getnframes()
    dur = (n / float(fr)) if fr else 0.0
    return ("# Audio: %s\n\n- **Channels:** %d\n- **Sample width:** %d-bit\n"
            "- **Sample rate:** %d Hz\n- **Frames:** %d\n- **Duration:** %.2f s\n"
            % (name, ch, sw * 8, fr, n, dur))


def media_note_fallback(path):
    """Generic, always-succeeds note so media output is never empty offline."""
    name = os.path.basename(path)
    try:
        size = os.path.getsize(path) if os.path.isfile(path) else 0
    except OSError:
        size = 0
    ext = os.path.splitext(path)[1].lstrip(".").upper() or "FILE"
    return ("# %s: %s\n\n- **Type:** %s\n- **Size:** %d bytes\n\n"
            "_Full extraction (transcription/OCR/metadata) for this format needs "
            "`markitdown[all]` plus its optional backend installed._\n"
            % (ext, name, ext, size))


IMAGE_EXTS = {".png", ".jpg", ".jpeg", ".gif", ".bmp", ".tiff", ".tif", ".webp", ".ppm"}
AUDIO_META_EXTS = {".mp3", ".m4a", ".flac", ".ogg", ".aac", ".wma"}

ERROR_FALLBACKS = {".docx": docx_fallback, ".xlsx": xlsx_fallback}

# Formats where a local extractor is MORE faithful than MarkItDown's native
# backend, used as the PRIMARY path (native remains the fallback).
#  - .xlsx: MarkItDown reads it via pandas, which coerces literal cell text like
#    "NA", "N/A", "NULL", "None" into missing values (NaN) - corrupting e.g. a
#    region code "NA". openpyxl preserves the literal stored values.
#  - .csv/.tsv: MarkItDown's CSV converter does not escape '|' or embedded
#    newlines, so a cell with a pipe adds a phantom column and a multiline cell
#    splits the row (a broken table). csv_fallback escapes cells; it also tabulates
#    .tsv (native passes TSV through as raw text) and sniffs ';' delimited CSVs.
PRIMARY_OVERRIDES = {
    ".xlsx": xlsx_fallback,
    ".csv": csv_fallback,
    ".tsv": lambda p: csv_fallback(p, delimiter="\t"),
}


def _empty_fallback_for(ext):
    if ext in IMAGE_EXTS:
        return image_fallback
    if ext == ".wav":
        return wav_fallback
    if ext in AUDIO_META_EXTS:
        return media_note_fallback
    return None


def _media_fallback_text(inp, ext):
    """Best-effort metadata for a media file (image/audio). Guaranteed non-empty
    for known media types - if the specific extractor itself fails, a generic
    note is used - so a media conversion is never lost to an empty file or a raw
    error. Returns None for non-media extensions."""
    fb = _empty_fallback_for(ext)
    if fb is None:
        return None
    try:
        out = fb(inp)
    except Exception:
        out = media_note_fallback(inp)  # always-succeeds guarantee for media
    return out if out.strip() else None


def convert_one(md, inp):
    """Convert via MarkItDown; fall back gracefully on errors or empty output."""
    stream_info = None
    ext = "" if is_url(inp) else os.path.splitext(inp)[1].lower()

    # Prefer a more faithful local extractor for select formats; on any problem
    # fall through to native MarkItDown so nothing is lost.
    if not is_url(inp) and os.path.isfile(inp):
        primary = PRIMARY_OVERRIDES.get(ext)
        if primary:
            try:
                out = primary(inp)
                if out.strip():
                    return out
            except Exception:
                pass

    if not is_url(inp) and os.path.isfile(inp) and ext in TEXT_EXTS:
        stream_info = _text_stream_info(inp, ext)

    try:
        result = md.convert(inp, stream_info=stream_info) if stream_info else md.convert(inp)
        text = getattr(result, "markdown", None) or getattr(result, "text_content", "") or ""
    except Exception as primary_error:
        if not is_url(inp) and os.path.isfile(inp):
            fb = ERROR_FALLBACKS.get(ext)
            if fb:
                try:
                    out = fb(inp)
                    sys.stderr.write("NOTE: used fallback extractor for %s "
                                     "(install 'markitdown[all]' for native fidelity)\n" % inp)
                    return out
                except Exception:
                    pass
            # Media files: a native backend is present but errored (e.g. no ffmpeg,
            # non-speech audio raising in the transcriber, a missing OCR model).
            # Degrade to metadata instead of failing, matching the empty-output path.
            media = _media_fallback_text(inp, ext)
            if media:
                sys.stderr.write("NOTE: used metadata fallback for %s "
                                 "(install 'markitdown[all]' + backend for full extraction)\n" % inp)
                return media
        raise primary_error

    # Enrichment: native succeeded but produced nothing (no exiftool/transcription).
    if (not text.strip()) and not is_url(inp) and os.path.isfile(inp):
        media = _media_fallback_text(inp, ext)
        if media:
            sys.stderr.write("NOTE: used metadata fallback for %s "
                             "(install 'markitdown[all]' + backend for full extraction)\n" % inp)
            return media
    return text


def resolve_output(multiple, output):
    """Decide (out_dir, out_file) from -o, validating cleanly. Returns a tuple;
    on an invalid combination prints a clear message and exits 1 (never a traceback)."""
    out_dir = None
    out_file = None
    try:
        if multiple:
            target_dir = output or "."
            if os.path.exists(target_dir) and not os.path.isdir(target_dir):
                sys.stderr.write("ERROR: -o must be a directory when multiple inputs are given "
                                 "(got existing file: %s)\n" % target_dir)
                sys.exit(1)
            os.makedirs(target_dir, exist_ok=True)
            out_dir = target_dir
        elif output:
            if output.endswith(("/", os.sep)) or os.path.isdir(output):
                if os.path.exists(output) and not os.path.isdir(output):
                    sys.stderr.write("ERROR: -o looks like a directory but a file exists there: %s\n" % output)
                    sys.exit(1)
                os.makedirs(output, exist_ok=True)
                out_dir = output
            else:
                parent = os.path.dirname(output)
                if parent:
                    os.makedirs(parent, exist_ok=True)
                out_file = output
    except OSError as e:
        sys.stderr.write("ERROR: could not prepare output location: %s\n" % e)
        sys.exit(1)
    return out_dir, out_file


def _stdout_pipe_closed():
    """The stdout consumer closed the pipe (e.g. `convert.py f | head`). The .md
    file is already written and is the authoritative result, so a closed echo pipe
    must not fail the run. Redirect stdout to devnull so the interpreter's
    shutdown flush can't re-raise BrokenPipeError after we return."""
    try:
        fd = sys.stdout.fileno()
    except Exception:
        return
    try:
        devnull = os.open(os.devnull, os.O_WRONLY)
        try:
            os.dup2(devnull, fd)
        finally:
            os.close(devnull)
    except Exception:
        pass


def _echo_stdout(text):
    """Echo Markdown to stdout (single-input / no -o case) without crashing on a
    Windows legacy-codepage console OR a closed downstream pipe. The .md file is
    already written as UTF-8; an echo problem must never turn a fully successful
    conversion into a false failure. On a console codec error (CJK/emoji outside
    the active code page) write faithful UTF-8 bytes to the binary buffer instead;
    on a closed pipe (the common `convert.py f | head`/`| less` preview workflow,
    which raises BrokenPipeError, a subclass of OSError - NOT UnicodeEncodeError)
    stop cleanly."""
    try:
        sys.stdout.write(text)
        return
    except UnicodeEncodeError:
        pass
    except (BrokenPipeError, OSError):
        _stdout_pipe_closed()
        return
    data = text.encode("utf-8", errors="replace")
    buf = getattr(sys.stdout, "buffer", None)
    if buf is not None:
        try:
            buf.write(data)
            buf.flush()
        except (BrokenPipeError, OSError):
            _stdout_pipe_closed()
    else:  # no binary buffer (rare wrapped stream): last-resort lossy echo
        enc = getattr(sys.stdout, "encoding", None) or "utf-8"
        try:
            sys.stdout.write(data.decode(enc, errors="replace"))
        except (BrokenPipeError, OSError):
            _stdout_pipe_closed()


def main():
    parser = argparse.ArgumentParser(
        description="Convert documents/URLs to Markdown via MarkItDown. Supports: " + SUPPORTED)
    parser.add_argument("inputs", nargs="+", help="File path(s), directory, or URL(s).")
    parser.add_argument("-o", "--output", help="Output .md file (single input) or directory.")
    parser.add_argument("--recursive", action="store_true", help="Recurse into sub-directories.")
    parser.add_argument("--use-plugins", action="store_true", help="Enable 3rd-party MarkItDown plugins.")
    parser.add_argument("--llm-model", help="LLM model for image descriptions (needs openai).")
    parser.add_argument("--llm-api-key", help="API key for the LLM client.")
    parser.add_argument("--llm-base-url", help="Base URL for an OpenAI-compatible endpoint.")
    parser.add_argument("--docintel-endpoint", help="Azure Document Intelligence endpoint.")
    args = parser.parse_args()

    ensure_markitdown()
    warn_if_old_markitdown()
    md = build_converter(args)

    inputs = expand_inputs(args.inputs, recursive=args.recursive)
    if not inputs:
        sys.stderr.write("No inputs to convert.\n")
        sys.exit(1)

    multiple = len(inputs) > 1
    out_dir, out_file = resolve_output(multiple, args.output)

    used_targets = set()
    errors = 0
    for inp in inputs:
        try:
            text = convert_one(md, inp)
        except Exception as e:
            sys.stderr.write("FAILED: %s -> %s\n" % (inp, e))
            errors += 1
            continue

        if out_file:
            target = out_file
        elif out_dir:
            target = unique_target(out_dir, out_name_for(inp, keep_ext=multiple), used_targets)
        else:
            target = out_name_for(inp)

        try:
            with open(target, "w", encoding="utf-8") as f:
                f.write(text)
        except OSError as e:
            sys.stderr.write("FAILED: could not write %s -> %s\n" % (target, e))
            errors += 1
            continue
        sys.stderr.write("OK: %s -> %s\n" % (inp, target))

        if not multiple and not out_file and not out_dir:
            _echo_stdout(text)

    if errors:
        sys.exit(2)


if __name__ == "__main__":
    try:
        main()
    except SystemExit:
        raise
    except KeyboardInterrupt:
        sys.stderr.write("\nInterrupted.\n")
        sys.exit(130)
    except Exception as e:
        sys.stderr.write("ERROR: %s\n" % e)
        sys.exit(1)
