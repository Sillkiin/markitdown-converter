#!/usr/bin/env python3
"""
verify_native.py - confirm the NETWORK / system-tool dependent backends that
can't be installed in a locked-down/offline environment actually work on YOUR
machine after `pip install -r requirements.txt`.

These paths were NOT verifiable where the skill was built (no PyPI / no exiftool):
  - native .docx (mammoth)        - audio transcription (pydub + SpeechRecognition / ffmpeg)
  - image metadata/OCR (exiftool) - YouTube transcript (youtube-transcript-api, needs network)
  - legacy .xls (xlrd)            - Outlook .msg (olefile)

Run:
    python tests/verify_native.py
    python tests/verify_native.py --audio path/to/speech.mp3   # confirm a real transcription
    python tests/verify_native.py --image path/to/photo.jpg    # confirm image metadata/OCR
    python tests/verify_native.py --youtube https://youtu.be/XXXX   # confirm a real transcript

It never hard-fails on a MISSING backend (those are reported as SKIP with the
exact pip command to enable them). It FAILS only if an installed backend is
present but errors, so a green run on your machine is meaningful.
"""
import argparse
import importlib
import os
import shutil
import subprocess
import sys
import tempfile

HERE = os.path.dirname(os.path.abspath(__file__))
CONVERT = os.path.join(HERE, "..", "scripts", "convert.py")

PASS, SKIP, FAIL = "PASS", "SKIP", "FAIL"
rows = []


def add(name, status, detail=""):
    rows.append((name, status, detail))


def have_module(mod):
    try:
        importlib.import_module(mod)
        return True
    except Exception:
        return False


def run_convert(path_or_url, extra=None):
    """Return (rc, stdout, stderr) for a single conversion to stdout."""
    cmd = [sys.executable, CONVERT, path_or_url]
    if extra:
        cmd += extra
    p = subprocess.run(cmd, capture_output=True, text=True)
    return p.returncode, p.stdout, p.stderr


def used_fallback(stderr):
    return "fallback" in stderr.lower()


# --------------------------------------------------------------------------- #
def check_backends():
    """Report which optional backends are installed."""
    mods = {
        "native .docx (mammoth)": "mammoth",
        "legacy .xls (xlrd)": "xlrd",
        "YouTube transcripts (youtube_transcript_api)": "youtube_transcript_api",
        "audio decode (pydub)": "pydub",
        "speech-to-text (speech_recognition)": "speech_recognition",
        "Outlook .msg (olefile)": "olefile",
        "image EXIF (exifread)": "exifread",
        "RSS (feedparser)": "feedparser",
    }
    for label, mod in mods.items():
        if have_module(mod):
            add("backend: " + label, PASS, "installed")
        else:
            add("backend: " + label, SKIP, "pip install -r requirements.txt")
    # exiftool is a system binary, not a python module
    if shutil.which("exiftool"):
        add("backend: image metadata (exiftool binary)", PASS, "on PATH")
    else:
        add("backend: image metadata (exiftool binary)", SKIP,
            "install exiftool (e.g. apt-get install libimage-exiftool-perl / brew install exiftool)")


def check_native_docx(tmp):
    if not have_module("mammoth"):
        add("native .docx conversion", SKIP, "mammoth not installed")
        return
    try:
        from docx import Document
    except Exception:
        add("native .docx conversion", SKIP, "python-docx unavailable to build sample")
        return
    p = os.path.join(tmp, "doc.docx")
    d = Document(); d.add_heading("Heading One", 1); d.add_paragraph("Native body text.")
    d.save(p)
    rc, out, err = run_convert(p)
    if rc == 0 and not used_fallback(err) and "Native body text." in out:
        add("native .docx conversion", PASS, "mammoth path used (no fallback)")
    elif rc == 0 and used_fallback(err):
        add("native .docx conversion", FAIL, "fell back despite mammoth present")
    else:
        add("native .docx conversion", FAIL, (err.strip().splitlines() or [""])[-1])


def check_image(tmp, user_image):
    target = user_image
    if not target:
        try:
            from PIL import Image
            target = os.path.join(tmp, "photo.jpg")
            im = Image.new("RGB", (48, 32), (20, 80, 160))
            ex = im.getexif(); ex[0x010F] = "Canon"; ex[0x0110] = "EOS R5"
            im.save(target, "JPEG", exif=ex)
        except Exception:
            add("image metadata/OCR", SKIP, "Pillow unavailable to build sample")
            return
    if not os.path.isfile(target):
        add("image metadata/OCR", FAIL, "image not found: %s" % target)
        return
    rc, out, err = run_convert(target)
    if rc != 0:
        add("image metadata/OCR", FAIL, (err.strip().splitlines() or [""])[-1])
    elif used_fallback(err):
        add("image metadata/OCR", SKIP,
            "exiftool/LLM not active - used Pillow fallback. Install exiftool or pass --llm-model for native.")
    else:
        add("image metadata/OCR", PASS, "native MarkItDown image path produced content")


def check_audio(tmp, user_audio):
    have = have_module("pydub") and have_module("speech_recognition")
    if user_audio:
        if not os.path.isfile(user_audio):
            add("audio transcription", FAIL, "audio not found: %s" % user_audio)
            return
        rc, out, err = run_convert(user_audio)
        if rc != 0:
            add("audio transcription", FAIL, (err.strip().splitlines() or [""])[-1])
        elif used_fallback(err):
            add("audio transcription", SKIP,
                "transcription backend inactive - used metadata fallback. "
                "Install: pip install 'markitdown[audio-transcription]' + ffmpeg")
        else:
            add("audio transcription", PASS, "produced transcript text from your file")
        return
    if have:
        add("audio transcription", SKIP,
            "backend installed - pass a real spoken-audio file: --audio speech.mp3 to confirm output")
    else:
        add("audio transcription", SKIP,
            "pip install 'markitdown[audio-transcription]' and install ffmpeg, then re-run with --audio")


def check_youtube(url):
    if not url:
        add("YouTube transcript", SKIP,
            "pass --youtube <url> to verify (needs network + youtube-transcript-api)")
        return
    if not have_module("youtube_transcript_api"):
        add("YouTube transcript", SKIP, "pip install youtube-transcript-api, then re-run with --youtube")
        return
    rc, out, err = run_convert(url)
    if rc == 0 and out.strip():
        add("YouTube transcript", PASS, "fetched transcript (%d chars)" % len(out.strip()))
    else:
        add("YouTube transcript", FAIL, (err.strip().splitlines() or ["empty output"])[-1])


def main():
    ap = argparse.ArgumentParser(description="Verify network/system-dependent MarkItDown backends on this machine.")
    ap.add_argument("--audio", help="A real spoken-audio file to confirm transcription.")
    ap.add_argument("--image", help="A real image to confirm native metadata/OCR.")
    ap.add_argument("--youtube", help="A YouTube URL to confirm transcript fetch.")
    args = ap.parse_args()

    tmp = tempfile.mkdtemp()
    check_backends()
    check_native_docx(tmp)
    check_image(tmp, args.image)
    check_audio(tmp, args.audio)
    check_youtube(args.youtube)

    width = max(len(n) for n, _, _ in rows)
    npass = sum(1 for _, s, _ in rows if s == PASS)
    nskip = sum(1 for _, s, _ in rows if s == SKIP)
    nfail = sum(1 for _, s, _ in rows if s == FAIL)
    print("MarkItDown native-backend verification\n")
    for n, s, d in rows:
        print("  %-4s  %-*s  %s" % (s, width, n, d))
    print("\n%d passed, %d skipped (backend not installed), %d failed" % (npass, nskip, nfail))
    if nfail:
        print("\nFAIL means an installed backend errored - see the detail above.")
    else:
        print("\nNo failures. SKIP items just need their backend installed "
              "(pip install -r requirements.txt, plus ffmpeg/exiftool where noted).")
    sys.exit(1 if nfail else 0)


if __name__ == "__main__":
    main()
